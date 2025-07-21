import docx
import pdfplumber
import io

def extract_text_from_pdf(file_input):
    """
    Extract text from a PDF file.

    Parameters:
        file_input (str or io.BytesIO): Path to the file or a file-like object.

    Returns:
        str: Extracted text from the PDF, or an error message.
    """
    text = ""
    try:
        if not isinstance(file_input, (str, io.BytesIO)):
            raise ValueError(f"Unsupported input type for PDF: {type(file_input)}")

        with pdfplumber.open(file_input) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"

        return text.strip()

    except Exception as e:
        return f"Error reading PDF: {e}"


def extract_text_from_docx(file_input):
    """
    Extract text from a DOCX file.

    Parameters:
        file_input (str or io.BytesIO): Path to the file or a file-like object.

    Returns:
        str: Extracted text from the DOCX, or an error message.
    """
    try:
        if isinstance(file_input, io.BytesIO):
            document = docx.Document(file_input)
        elif isinstance(file_input, str):
            with open(file_input, 'rb') as f:
                document = docx.Document(f)
        else:
            raise ValueError(f"Unsupported input type for DOCX: {type(file_input)}")

        return '\n'.join([para.text for para in document.paragraphs if para.text.strip()])

    except Exception as e:
        return f"Error reading DOCX: {e}"
